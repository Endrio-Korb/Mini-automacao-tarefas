
import os
import csv

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.shared import Inches


if __name__ == "__main__":

    document = Document()

    # Estilos da informações
    styles = document.styles
    p = styles.add_style("paragrafo", WD_STYLE_TYPE.PARAGRAPH) 
    p.font.name = "Calibri Light (Títulos)"
    p.font.size = Pt(45)
    p.font.bold = True

    with open(os.path.join(os.getcwd(), "Organização das salas.csv"), "r") as arquivo:
        arquivo_csv = csv.reader(arquivo, delimiter = ";")


        for linha in arquivo_csv:
            sala = "SALA " + str(linha[1])
            projetor = "PROJETOR " + str(linha[2])
            patrimonio = "PATRIMÔNIO " + str(linha[3])
            capacidade = "CAPACIDADE " +str(linha[4])

            paragraph = document.add_paragraph(sala, style="paragrafo")
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraph = document.add_paragraph(projetor,style="paragrafo")
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraph = document.add_paragraph(patrimonio, style="paragrafo")
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraph = document.add_paragraph(capacidade, style="paragrafo")
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraph = document.add_picture("logo.png", width=Inches(6.5))
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            document.save(f"{sala}.docx")
