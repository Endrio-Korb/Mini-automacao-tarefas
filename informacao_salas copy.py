
import os
import csv
import win32print
import win32api
from docx import Document
import pandas as pd
import openpyxl
import numpy



if __name__ == "__main__":

    tabela = pd.read_excel("Organização das salas.xlsx")
    documento = Document("modelo_informacao.docx")

    for linha in tabela.index:
        for paragrafo in documento.paragraphs:

            sala = str(tabela.loc[linha, "Sala"])
            projetor = str(tabela.loc[linha, "Modelo Projetor"])
            patrimonio = str(tabela.loc[linha, "Patromonio"])
            capacidade = str(tabela.loc[linha, "Qtd Cadeiras"])

            referencia = {
                "XXXX" : sala ,
                "XXXXXXXXX" : projetor,
                "XXXXXXXXXXX" : patrimonio,
                "XX" : capacidade,
            }
        
            for paragrafo in documento.paragraphs:
                for codigo in referencia:
                    valor = referencia[codigo]
                    paragrafo.text = paragrafo.text.replace(codigo, valor)

            documento.save(f"Sala - {sala}.docx")

                

    
